#!/usr/bin/env python3
"""Genera las instrucciones de reproducción en formato .docx, en español.

El texto vive aquí y no en un archivo suelto para que las rutas y los comandos
salgan del mismo lugar que el paquete, y no se desincronicen como ya pasó con
el script de calibración que un revisor recibió desactualizado.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT.parent / "Instrucciones_Reproduccion_EM_Audio.docx"

MONO = "Consolas"


def code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = MONO
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    doc.add_paragraph()


def main() -> int:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    t = doc.add_heading("Reproducción independiente de EM-Audio", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run("Para: ").bold = True
    p.add_run("Daniel A. Balderrama-Alvarez, Universidad de Sonora "
              "(ORCID 0009-0002-5180-0406)")

    doc.add_paragraph(
        "El paquete adjunto trae el repositorio completo en la versión "
        "etiquetada v1.0.0. No hace falta clonar nada ni tener conexión de red "
        "durante la ejecución, salvo la descarga del corpus de audio, que el "
        "propio script hace la primera vez.")

    doc.add_heading("Qué ejecutar", level=1)
    code(doc, [
        "unzip EM_Audio_reproduction_package.zip",
        "cd EM_Audio_reproduction_package/em-audio",
        "pip install -r requirements.txt",
        "./run_all.sh",
        "python3 tools/verify_reproduction.py",
    ])
    doc.add_paragraph(
        "Requisitos de línea de comandos: FFmpeg, ffprobe, Node, eSpeak NG y "
        "c2patool en el PATH. Los paquetes de Python están en "
        "requirements.txt, que es el paso que faltaba en la versión anterior "
        "de estas instrucciones: matplotlib y numpy para las figuras, y "
        "piper-tts para el brazo de robustez con TTS neuronal.")
    doc.add_paragraph(
        "run_all.sh ahora revisa las dependencias antes de empezar y se "
        "detiene con código 2 si falta alguna, en lugar de fallar a la mitad. "
        "La corrida completa tarda alrededor de veinte minutos, más la "
        "descarga del corpus.")
    doc.add_paragraph(
        "Las versiones exactas con las que se produjeron los resultados están "
        "en results/PREFLIGHT.txt. Las tuyas van a ser distintas y eso es "
        "parte del punto del ejercicio.")

    doc.add_heading("Qué compara la herramienta", level=1)
    doc.add_paragraph(
        "verify_reproduction.py aplica la separación que el manuscrito "
        "describe, en lugar de dejarla a criterio:")
    for head, body in (
        ("Salidas deterministas",
         "deben coincidir exactamente. Conteos de conformidad y totales de "
         "aprobado/fallado, clasificaciones de promoción, conjuntos de "
         "linaje, estructura de intervalos, comportamiento del canal de "
         "soporte, batería de scope, acuerdo del oráculo y digests de esencia "
         "decodificada. Cualquier diferencia aquí devuelve exit distinto de "
         "cero."),
        ("Salidas dependientes del entorno",
         "se espera que difieran y no hacen fallar la corrida. Tiempos de "
         "reloj, su dispersión, y los bytes de un archivo recién firmado, que "
         "cambian porque la firma incorpora una marca de tiempo."),
    ):
        b = doc.add_paragraph(style="List Bullet")
        b.add_run(f"{head}: ").bold = True
        b.add_run(body)
    doc.add_paragraph(
        "Una corrida en la que los tiempos coincidieran exactamente sería el "
        "resultado sospechoso, no el tranquilizador.")

    doc.add_heading("Qué necesito de vuelta", level=1)
    for item in (
        "La salida completa de run_all.sh y de verify_reproduction.py.",
        "Tu results/PREFLIGHT.txt, que registra tu máquina y las versiones de "
        "las herramientas.",
        "La carpeta results/machine_readable/ completa.",
        "Sistema operativo, hardware y fecha.",
    ):
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Si algo no coincide", level=1)
    doc.add_paragraph(
        "No lo suavices y no asumas que es tu entorno. Mándalo tal cual y lo "
        "clasificamos juntos: diferencia de entorno, diferencia de versión de "
        "herramienta, ambigüedad en cómo está especificado el contrato, o "
        "discrepancia genuina.")
    p = doc.add_paragraph()
    p.add_run(
        "Una reproducción que falla y se reporta con honestidad vale más para "
        "este trabajo que una que pasa en silencio.").bold = True

    doc.add_heading("Un caso concreto que puede aparecer", level=1)
    doc.add_paragraph(
        "El experimento de contención de soporte mide hasta dónde alcanza la "
        "dependencia de cada operador y la compara contra el límite declarado. "
        "Ese límite se calibró con FFmpeg 9.0.1 en macOS. Con otra versión de "
        "FFmpeg, el codificador MP3 puede tener un alcance distinto, y si "
        "resulta mayor que lo declarado el experimento va a fallar y "
        "run_all.sh va a salir con error.")
    doc.add_paragraph(
        "Si eso pasa, no es un error tuyo ni algo que haya que arreglar antes "
        "de reportarlo. Es exactamente lo que el manuscrito advierte que puede "
        "ocurrir, y confirmarlo desde fuera vale más que evitarlo. Mándame el "
        "resultado tal como salga.")

    doc.save(OUT)
    print(f"[docx] {OUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
