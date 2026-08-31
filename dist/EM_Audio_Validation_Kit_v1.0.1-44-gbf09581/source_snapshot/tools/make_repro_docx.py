#!/usr/bin/env python3
"""Genera las instrucciones de reproducción en formato .docx, en español.

El texto vive aquí y no en un archivo suelto para que las rutas y los comandos
salgan del mismo lugar que el paquete, y no se desincronicen como ya pasó con
el script de calibración que un revisor recibió desactualizado.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT.parent / "Instrucciones_Reproduccion_EM_Audio.docx"

MONO = "Consolas"


def code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = MONO
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    doc.add_paragraph()


def main() -> int:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    t = doc.add_heading("Reproducción independiente de EM-Audio", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run("Para: ").bold = True
    p.add_run("Daniel A. Balderrama-Alvarez, Universidad de Sonora "
              "(ORCID 0009-0002-5180-0406)")

    doc.add_paragraph(
        "El paquete adjunto trae el repositorio completo en la versión "
        "etiquetada v1.0.1. No hace falta clonar nada. La única descarga es el "
        "corpus de audio, y el propio script la hace la primera vez.")

    doc.add_heading("Qué ejecutar", level=1)
    code(doc, [
        "unzip EM_Audio_reproduction_package.zip",
        "cd EM_Audio_reproduction_package/em-audio",
        "pip install -r requirements.txt",
        "./run_all.sh",
        "python3 tools/verify_reproduction.py",
    ])
    doc.add_paragraph(
        "En el PATH necesitas FFmpeg, ffprobe, Node, eSpeak NG y c2patool. Los "
        "paquetes de Python están en requirements.txt: matplotlib y numpy para "
        "las figuras, piper-tts para el brazo de robustez. Ese paso de pip "
        "faltaba en las instrucciones que te mandé antes, y es la razón de que "
        "tu primera corrida se quedara sin esos dos módulos.")
    doc.add_paragraph(
        "run_all.sh ahora revisa las dependencias antes de empezar. Si falta "
        "alguna se detiene con código 2 y te dice cuál, en lugar de fallar a "
        "los veinte minutos. La corrida completa tarda alrededor de eso, más "
        "la descarga del corpus.")
    doc.add_paragraph(
        "Las versiones con las que se produjeron los resultados de referencia "
        "están en results/PREFLIGHT.txt. Las tuyas serán distintas, y esa "
        "diferencia es lo que estamos midiendo.")

    doc.add_heading("Qué compara la herramienta", level=1)
    doc.add_paragraph(
        "verify_reproduction.py separa las salidas en dos clases, siguiendo lo "
        "que dice el manuscrito, para que la clasificación no dependa del "
        "criterio de quien mira el resultado.")
    doc.add_paragraph(
        "Las deterministas tienen que coincidir exactamente, y cualquier "
        "diferencia devuelve exit distinto de cero:")
    for item in ("conteos de conformidad y totales de aprobado y fallado",
                 "clasificaciones de promoción y conjuntos de linaje",
                 "estructura de intervalos y comportamiento del canal de soporte",
                 "batería de scope y acuerdo del oráculo",
                 "digests de esencia decodificada"):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "Las dependientes del entorno van a diferir y no hacen fallar la "
        "corrida: tiempos de reloj, su dispersión, y los bytes de un archivo "
        "recién firmado, que cambian porque la firma lleva una marca de "
        "tiempo. Si tus tiempos salieran idénticos a los míos, habría que "
        "desconfiar de la medición.")

    doc.add_heading("Qué necesito de vuelta", level=1)
    for item in (
        "La salida completa de run_all.sh y de verify_reproduction.py.",
        "Tu results/PREFLIGHT.txt, que registra tu máquina y las versiones.",
        "La carpeta results/machine_readable/ completa.",
        "Sistema operativo, hardware y fecha.",
    ):
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph(
        "El directorio results/machine_readable/ va vacío a propósito. En el "
        "paquete anterior iba con mis resultados dentro, así que los "
        "experimentos que fallaron dejaron mis archivos en su sitio y la "
        "comparación los dio por coincidentes. Tu C2_robustness.json de la "
        "primera corrida era mi archivo, no el tuyo.")

    doc.add_heading("Si algo no coincide", level=1)
    doc.add_paragraph(
        "No lo suavices y no asumas que es tu entorno. Mándalo tal cual y lo "
        "clasificamos juntos: diferencia de entorno, diferencia de versión de "
        "herramienta, ambigüedad en cómo está especificado el contrato, o "
        "discrepancia real del contrato. Un fallo reportado sin ajustar es "
        "más útil aquí que una corrida que pasa.")

    doc.add_heading("El caso MP3, que ya apareció", level=1)
    doc.add_paragraph(
        "El experimento de contención mide hasta dónde alcanza la dependencia "
        "de cada operador y lo compara contra el límite declarado. Ese límite "
        "se calibró con FFmpeg 9.0.1 en macOS. Tu FFmpeg 8.0.1 midió un "
        "alcance mayor para el codificador MP3, así que el experimento falló y "
        "run_all.sh salió con error.")
    doc.add_paragraph(
        "Eso no es un error tuyo. Es lo que el manuscrito advierte que puede "
        "pasar cuando cambia el build, y confirmarlo desde fuera es "
        "precisamente lo que hacía falta. Los otros seis operadores "
        "reprodujeron su alcance exactamente, cruzando arquitecturas "
        "distintas.")
    doc.add_paragraph(
        "En esta segunda corrida espero los mismos dos números que la primera. "
        "Si salen distintos, eso sí sería un problema, porque significaría que "
        "la medición no es reproducible ni consigo misma.")

    doc.save(OUT)
    print(f"[docx] {OUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
